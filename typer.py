import asyncio
from breaker import RealKeyboardTyper
from models import DocumentContent, RetypedDocument
from pydantic_ai import Agent
from dotenv import load_dotenv
import traceback
from pynput.keyboard import Key, Controller as KeyboardController
from pynput.mouse import Button, Controller as MouseController



load_dotenv()




task_desp = open("system_prompt.txt", "r").read()

class DocumentRetyper:
    """Class to retype documents with real keyboard typing"""
    def __init__(self, delay=0.005,model_name: str = 'google-gla:gemini-2.0-flash'):
        self.model_name = model_name
        self.document_retyper = None
        self.keyboard_typer = RealKeyboardTyper(delay=delay)  # Use the configurable delay
        
    async def async_init(self):
        """Async initialization method"""
        self.document_retyper = Agent(
            model=self.model_name,
            # 'groq:llama-3.3-70b-versatile',
            result_type=RetypedDocument,
            system_prompt= task_desp
        )
        return self

    async def retype_document_with_real_typing(self, document_text: str, typing_position=None, error_correction=True, progress_callback=None):
        """Function to retype a document using real keyboard inputs with formatting support"""
        content = DocumentContent(text=document_text)
        
        # Just show a simple message in terminal
        print(f"Reading document content...")
        
        # Run the agent silently
        result = await self.document_retyper.run(
            f"Please retype the following document exactly as it is, preserving all formatting, spacing,"
            f" line breaks, and paragraph structure: {content.text}"
        )
        
        # Perform real keyboard typing with formatting support
        try:
            typed_chars = 0
        
        # Function to track progress during typing
            async def progress_tracker(chars_typed):
                nonlocal typed_chars
                typed_chars = chars_typed
                if progress_callback:
                    await progress_callback(typed_chars)
        
            # Check if the content likely has formatting markers
            if "**" in result.data.content or "__" in result.data.content or "_" in result.data.content:
                # Use the enhanced typing with formatting method
                await self.keyboard_typer.type_with_formatting(result.data.content, typing_position,progress_tracker)
            else:
                # Use the regular typing with verification method
                await self.keyboard_typer.type_with_verification(result.data.content, typing_position, error_correction, progress_tracker)
            print(f"\nDocument successfully retyped using real keyboard inputs!")
        except Exception as e:
            print(f"Error during typing: {str(e)}")
            traceback.print_exc()
        
        return result.data.content
    
    # function for CTK only 
    # In typer.py, within the DocumentRetyper class
    # async def retype_document_for_ctk(self, document_text: str, typing_position=None, error_correction=True):
    #     """Function to retype a document using real keyboard inputs for Customtkinter app, without progress tracking."""
    #     import logging
    #     logger = logging.getLogger(__name__)
        
    #     logger.info("Starting retype_document_for_ctk...")
    #     content = DocumentContent(text=document_text)
        
    #     # Just show a simple message in terminal
    #     print(f"Reading document content...")
    #     logger.info("Reading document content...")
        
    #     # Run the agent silently
    #     logger.info("Calling document_retyper.run...")
    #     result = await self.document_retyper.run(
    #         f"Please retype the following document exactly as it is, preserving all formatting, spacing,"
    #         f" line breaks, and paragraph structure: {content.text}"
    #     )
    #     logger.info("document_retyper.run completed.")
        
    #     # Perform real keyboard typing with formatting support
    #     try:
    #         logger.info("Checking for formatting markers...")
    #         # Check if the content likely has formatting markers
    #         if "**" in result.data.content or "__" in result.data.content or "_" in result.data.content:
    #             logger.info("Using type_with_formatting...")
    #             # Use the enhanced typing with formatting method, without progress_tracker
    #             await self.keyboard_typer.type_with_formatting(result.data.content, typing_position)
    #             logger.info("type_with_formatting completed.")
    #         else:
    #             logger.info("Using type_with_verification...")
    #             # Use the regular typing with verification method, without progress_tracker
    #             await self.keyboard_typer.type_with_verification(result.data.content, typing_position, error_correction)
    #             logger.info("type_with_verification completed.")
    #         print(f"\nDocument successfully retyped using real keyboard inputs!")
    #         logger.info("Document successfully retyped.")
    #     except Exception as e:
    #         print(f"Error during typing: {str(e)}")
    #         logger.error(f"Error during typing: {str(e)}")
    #         raise  # Re-raise the exception to be handled by the caller
        
    #     return result.data.content
# In typer.py, within the DocumentRetyper class
    async def retype_document_for_ctk(self, document_text: str, typing_position=None, error_correction=True):
        """Function to retype a document using real keyboard inputs for Customtkinter app, without progress tracking."""
        import logging
        import traceback
        logger = logging.getLogger(__name__)
        
        logger.info("Starting retype_document_for_ctk...")
        content = DocumentContent(text=document_text)
        
        # Just show a simple message in terminal
        print(f"Reading document content...")
        logger.info("Reading document content...")
        
        # Run the agent silently with a timeout
        logger.info("Calling document_retyper.run...")
        try:
            result = await asyncio.wait_for(
                self.document_retyper.run(
                    f"Please retype the following document exactly as it is, preserving all formatting, spacing,"
                    f" line breaks, and paragraph structure: {content.text}"
                ),
                timeout=30  # Timeout after 30 seconds
            )
            logger.info("document_retyper.run completed.")
            logger.debug(f"Result from document_retyper.run: {result.data.content[:100]}...")  # Log first 100 chars of result
        except asyncio.TimeoutError:
            logger.error("API call to document_retyper.run timed out after 30 seconds.")
            raise Exception("API call timed out. Please check your network connection and try again.")
        except Exception as e:
            logger.error(f"Error during document_retyper.run: {str(e)}")
            logger.debug(traceback.format_exc())
            raise Exception(f"Failed to process document with AI model: {str(e)}")
        
        # Perform real keyboard typing with formatting support
        try:
            logger.info("Checking for formatting markers...")
            # Check if the content likely has formatting markers
            if "**" in result.data.content or "__" in result.data.content or "_" in result.data.content:
                logger.info("Using type_with_formatting...")
                # Use the enhanced typing with formatting method, without progress_tracker
                await self.keyboard_typer.type_with_formatting(result.data.content, typing_position)
                logger.info("type_with_formatting completed.")
            else:
                logger.info("Using type_with_verification...")
                # Use the regular typing with verification method, without progress_tracker
                await self.keyboard_typer.type_with_verification(result.data.content, typing_position, error_correction)
                logger.info("type_with_verification completed.")
            print(f"\nDocument successfully retyped using real keyboard inputs!")
            logger.info("Document successfully retyped.")
        except Exception as e:
            print(f"Error during typing: {str(e)}")
            logger.error(f"Error during typing: {str(e)}")
            logger.debug(traceback.format_exc())
            raise  # Re-raise the exception to be handled by the caller
        
        return result.data.content

    async def display_document_info(self, file_path: str):
        """Display basic document info without typing the content"""
        # Check if it's a docx file
        if file_path.endswith('.docx'):
            import docx
            doc = docx.Document(file_path)
            # Count paragraphs and characters in docx
            paragraphs = len(doc.paragraphs)
            char_count = sum(len(paragraph.text) for paragraph in doc.paragraphs)
            
            print(f"Document loaded: {file_path}")
            print(f"Document contains {paragraphs} paragraphs and {char_count} characters.")
            print("Starting retyping process...\n")
            return f"Document contains {paragraphs} paragraphs and {char_count} characters."
        else:
            # For regular text files
            with open(file_path, 'r') as f:
                content = f.read()
                
            print(f"Document loaded: {file_path}")
            line_count = content.count('\n') + 1
            char_count = len(content)
            print(f"Document contains {line_count} lines and {char_count} characters.")
            print("Starting retyping process...\n")
            return f"Document contains {line_count} lines and {char_count} characters."

# Function to extract text from docx with better formatting preservation
def extract_text_from_docx(docx_path):
    """Extract text from docx file with enhanced formatting preservation"""
    import docx
    doc = docx.Document(docx_path)
    full_text = []
    
    # Process paragraphs with enhanced formatting
    for para in doc.paragraphs:
        # Don't skip empty paragraphs - they're part of formatting
        # Handle heading styles
        if para.style.name.startswith('Heading'):
            # Add appropriate formatting based on heading level
            level = int(para.style.name.replace('Heading', '')) if para.style.name != 'Heading' else 1
            prefix = '#' * level + ' ' if level > 0 else ''
            full_text.append(f"{prefix}{para.text}")
        else:
            # Regular paragraph - include even if empty to preserve spacing
            full_text.append(para.text)
    
    # Process tables if any
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            table_rows.append(" | ".join(row_text))
        full_text.append("\n".join(table_rows))
    
    # Join with paragraph spacing preserved
    return "\n\n".join(full_text)